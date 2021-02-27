import csv
import docx
from docx import Document

names = []
with open('export.csv') as csv_file:
    csv_reader = csv.reader(csv_file, delimiter=',')
    for row in csv_reader:
            names.append (row[1] + " " + row[2])

combined_document = Document('combined_document.docx')

for i in range (1,len(names)):
    original_document = Document('template.docx')
    for paragraph in original_document.paragraphs:
        if 'Name of Member of Parliament' in paragraph.text:
            mp_name = names [i]
            paragraph.text = paragraph.text.replace('Name of Member of Parliament', mp_name)
            original_document.save(mp_name + ".docx")
        combined_document.add_paragraph(paragraph.text)
        para = combined_document.add_paragraph ()
    combined_document.add_page_break()
    combined_document.save('combined_document.docx')
    

